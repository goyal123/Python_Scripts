import docx
from docx.shared import Inches

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Image in its native size
doc.add_heading('Image in Native Size:', 3)
doc.add_picture('a.png',width=Inches(7), height=Inches(4))
doc.add_heading('Geeks', 3).bold=True
doc.add_picture('b.png', width=Inches(7), height=Inches(4))

# Now save the document to a location
doc.save('Document.docx')
